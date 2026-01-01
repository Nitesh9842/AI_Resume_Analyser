"""
Text Extraction Module
Handles PDF, DOCX, DOC, and TXT file formats
"""

import pdfplumber
import fitz  # PyMuPDF
import docx2txt
from docx import Document
import re
from typing import Optional, BinaryIO, Union
import io


def extract_text_from_pdf_pdfplumber(file) -> str:
    """Extract text from PDF using pdfplumber"""
    try:
        text = ""
        # Handle both file path and file object
        if isinstance(file, bytes):
            pdf_file = io.BytesIO(file)
        elif isinstance(file, str):
            pdf_file = file
        else:
            pdf_file = io.BytesIO(file.read()) if hasattr(file, 'read') else file
            
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error with pdfplumber: {e}")
        return ""


def extract_text_from_pdf_pymupdf(file) -> str:
    """Extract text from PDF using PyMuPDF (fallback)"""
    try:
        text = ""
        # Read file content
        if hasattr(file, 'read'):
            file_content = file.read()
            if hasattr(file, 'seek'):
                file.seek(0)
        else:
            with open(file, 'rb') as f:
                file_content = f.read()
                
        pdf_document = fitz.Document(stream=file_content, filetype="pdf")
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text += page.get_text() + "\n" # type: ignore
        pdf_document.close()
        return text
    except Exception as e:
        print(f"Error with PyMuPDF: {e}")
        return ""


def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        # Convert to BytesIO if needed
        if hasattr(file, 'read'):
            file_obj = io.BytesIO(file.read())
            if hasattr(file, 'seek'):
                file.seek(0)
        else:
            file_obj = file
            
        # Method 1: Using docx2txt
        try:
            text = docx2txt.process(file_obj)
            if text.strip():
                return text
        except:
            pass
        
        # Method 2: Using python-docx (fallback)
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
        doc = Document(file_obj)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


def extract_text_from_txt(file) -> str:
    """Extract text from TXT file"""
    try:
        if hasattr(file, 'read'):
            content = file.read()
            if isinstance(content, bytes):
                try:
                    return content.decode('utf-8')
                except UnicodeDecodeError:
                    return content.decode('latin-1')
            return content
        else:
            with open(file, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        print(f"Error reading TXT file: {e}")
        return ""


def extract_text_from_file(file) -> Optional[str]:
    """
    Main function to extract text from uploaded file
    Supports: PDF, DOCX, DOC, TXT
    file can be: file path (str), file object, or werkzeug FileStorage
    """
    if file is None:
        return None
    
    # Get filename
    if hasattr(file, 'filename'):
        filename = file.filename
    elif hasattr(file, 'name'):
        filename = file.name
    elif isinstance(file, str):
        filename = file
    else:
        return None
        
    file_type = filename.split('.')[-1].lower()
    
    try:
        if file_type == 'pdf':
            # Try pdfplumber first
            text = extract_text_from_pdf_pdfplumber(file)
            if not text.strip():
                # Reset file pointer if possible
                if hasattr(file, 'seek') and not isinstance(file, str):
                    file.seek(0)
                # Fallback to PyMuPDF
                text = extract_text_from_pdf_pymupdf(file)
            return text
        
        elif file_type in ['docx', 'doc']:
            return extract_text_from_docx(file)
        
        elif file_type == 'txt':
            return extract_text_from_txt(file)
        
        else:
            print(f"Unsupported file format: {file_type}")
            return None
            
    except Exception as e:
        print(f"Error processing file: {e}")
        return None


def clean_text(text: str) -> str:
    """
    Clean and normalize text
    - Remove extra whitespace
    - Remove special characters (keep alphanumeric and basic punctuation)
    - Normalize line breaks
    """
    if not text:
        return ""
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep letters, numbers, and basic punctuation
    text = re.sub(r'[^\w\s\.\,\-\+\#\(\)\/\&]', '', text)
    
    # Normalize line breaks
    text = text.replace('\n', ' ').strip()
    
    return text


def extract_email(text: str) -> Optional[str]:
    """Extract email address from text"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-ZaZ0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from text"""
    phone_pattern = r'(\+?\d{1,3}[-.\s]??)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    return phones[0] if phones else None


def extract_urls(text: str) -> list:
    """Extract URLs from text"""
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    return re.findall(url_pattern, text)